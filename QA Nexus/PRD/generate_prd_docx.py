from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
CHARTS_DIR = BASE_DIR / "prd_charts"

ACCENT = RGBColor(18, 52, 86)
MUTED = RGBColor(90, 98, 108)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_spacing(paragraph, before=0, after=4, line=1.15) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_runs(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part.strip("`*"))
        if part.startswith("**") and part.endswith("**"):
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)


def parse_blocks(lines: list[str]) -> list[tuple[str, object]]:
    blocks: list[tuple[str, object]] = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            blocks.append(("hr", None))
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            blocks.append(("heading", (level, text)))
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = []
            for idx, row in enumerate(table_lines):
                if idx == 1 and set(row.replace("|", "").replace("-", "").replace(":", "").strip()) == set():
                    continue
                rows.append([cell.strip() for cell in row.strip("|").split("|")])
            if rows:
                blocks.append(("table", rows))
            continue

        if re.match(r"^\d+\.\s+", stripped):
            items = []
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                items.append(re.sub(r"^\d+\.\s+", "", lines[i].strip()))
                i += 1
            blocks.append(("numbered", items))
            continue

        if stripped.startswith("- "):
            items = []
            while i < len(lines) and lines[i].strip().startswith("- "):
                items.append(lines[i].strip()[2:].strip())
                i += 1
            blocks.append(("bullets", items))
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            next_line = lines[i].strip()
            if (
                not next_line
                or next_line == "---"
                or next_line.startswith("#")
                or next_line.startswith("|")
                or next_line.startswith("- ")
                or re.match(r"^\d+\.\s+", next_line)
            ):
                break
            paragraph_lines.append(next_line)
            i += 1
        blocks.append(("paragraph", " ".join(paragraph_lines)))

    return blocks


def add_cover(document: Document, metadata: dict[str, str], title: str, subtitle: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=40, after=10, line=1.0)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = ACCENT

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=24, line=1.0)
    run = p.add_run(subtitle)
    run.font.size = Pt(15)
    run.font.color.rgb = MUTED

    info = document.add_table(rows=5, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.style = "Table Grid"
    fields = [
        ("Version", metadata.get("Document Version", "v6.0")),
        ("Status", metadata.get("Document Status", "Draft for Leadership Review")),
        ("Last Updated", metadata.get("Last Updated", "2026-04-22")),
        ("Organization", metadata.get("Organization", "Iksula Services Pvt Ltd")),
        ("Release", metadata.get("Release", "MVP")),
    ]
    for row, (label, value) in zip(info.rows, fields):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], "EAF1F8")
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    document.add_paragraph()
    document.add_page_break()


def add_table(document: Document, rows: list[list[str]]) -> None:
    cols = max(len(r) for r in rows)
    table = document.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            value = row[c_idx] if c_idx < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, before=0, after=0, line=1.0)
            add_runs(p, value)
            if r_idx == 0:
                set_cell_shading(cell, "D9EAF7")
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = ACCENT
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    document.add_paragraph()


def maybe_add_image(document: Document, heading_text: str) -> None:
    mapping = {
        "Delivery Roadmap": "timeline_gantt.png",
        "Risks and Mitigations": "risk_heatmap.png",
        "Solution Overview": "architecture_layers.png",
    }
    filename = mapping.get(heading_text)
    if not filename:
        return
    path = CHARTS_DIR / filename
    if not path.exists():
        return
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(path), width=Inches(6.5))
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(caption, before=0, after=10, line=1.0)
    cap_run = caption.add_run(path.stem.replace("_", " ").title())
    cap_run.italic = True
    cap_run.font.size = Pt(9)
    cap_run.font.color.rgb = MUTED


def extract_metadata(text: str) -> dict[str, str]:
    metadata = {}
    for line in text.splitlines():
        if ":" in line and line.startswith("**") and line.endswith("**") is False:
            continue
    for key in ["Organization", "Document Version", "Document Status", "Last Updated", "Release"]:
        match = re.search(rf"\*\*{re.escape(key)}:\*\*\s*(.+)", text)
        if match:
            metadata[key] = match.group(1).strip()
    metadata.setdefault("Release", "MVP")
    return metadata


def extract_title_and_subtitle(lines: list[str]) -> tuple[str, str]:
    title = "QA Nexus"
    subtitle = "Product Requirements Document"
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("# "):
            title = stripped[2:].strip()
            break
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("## "):
            subtitle = stripped[3:].strip()
            break
    return title, subtitle


def build_docx(md_path: Path, docx_path: Path) -> None:
    md_text = md_path.read_text(encoding="utf-8")
    lines = md_text.splitlines()
    title, subtitle = extract_title_and_subtitle(lines)
    start_idx = 0
    for idx, line in enumerate(lines):
        if line.strip() == "## 1. Document Control":
            start_idx = idx
            break

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)

    for level, size in [(1, 18), (2, 14), (3, 11.5)]:
        style = styles[f"Heading {level}"]
        style.font.name = "Aptos"
        style.font.bold = True
        style.font.color.rgb = ACCENT
        style.font.size = Pt(size)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"{title} | {extract_metadata(md_text).get('Document Version', 'v1.0')}")
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED

    add_cover(document, extract_metadata(md_text), title, subtitle)

    blocks = parse_blocks(lines[start_idx:])
    skip_first_h1 = True
    for kind, payload in blocks:
        if kind == "heading":
            level, text = payload
            if skip_first_h1 and level == 1 and text == title:
                skip_first_h1 = False
                continue
            if level == 2 and text == subtitle:
                continue
            p = document.add_paragraph(style=f"Heading {min(level, 3)}")
            set_paragraph_spacing(p, before=8 if level > 1 else 0, after=4, line=1.0)
            p.add_run(text)
            maybe_add_image(document, text)
        elif kind == "paragraph":
            p = document.add_paragraph()
            set_paragraph_spacing(p, before=0, after=6, line=1.15)
            add_runs(p, payload)
        elif kind == "bullets":
            for item in payload:
                p = document.add_paragraph(style="List Bullet")
                set_paragraph_spacing(p, before=0, after=2, line=1.1)
                add_runs(p, item)
        elif kind == "numbered":
            for item in payload:
                p = document.add_paragraph(style="List Number")
                set_paragraph_spacing(p, before=0, after=2, line=1.1)
                add_runs(p, item)
        elif kind == "table":
            add_table(document, payload)
        elif kind == "hr":
            document.add_paragraph()

    document.save(docx_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate a DOCX PRD from a markdown file.")
    parser.add_argument("input_md", nargs="?", default=str(BASE_DIR / "PRD.md"))
    parser.add_argument("output_docx", nargs="?", default=str(BASE_DIR / "PRD.docx"))
    return parser.parse_args()


if __name__ == "__main__":
    args = parse_args()
    build_docx(Path(args.input_md), Path(args.output_docx))
